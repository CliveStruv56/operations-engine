"""DOCX assembly: cover page with a DRAFT header, headings, the module's own
data tables, endnote-style citations and a Data sources appendix.

Citation ids resolve only against the excerpts the pack supplied — fabricated
markers are stripped and counted, never rendered as fake references.
`[TO CONFIRM: …]` markers stay visible in the text and are counted for the
UI's "N items to confirm" panel. Both are product requirements, so they live
here rather than in any one module.
"""

import io
import re
from collections.abc import Callable
from dataclasses import dataclass, field
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from worker.drafting.pack import DraftPackBase, VaultExcerpt
from worker.drafting.sections import Section, measure

# Two shapes, because the `c:` prefix decides how much benefit of the doubt a
# marker gets.
#
# Prefixed: the model is telling us it meant a citation, so whatever follows is
# a citation attempt whatever it looks like, and an unresolvable one is a
# hallucination to strip. Requiring hex here let `[c:s1]` — a wholly invented
# id — through to a funder's form, because `s` is not a hex digit and the
# pattern simply did not match (seen live 11 Aug 2026).
#
# Prefix-less: nothing distinguishes it from prose, so it must still look like
# a truncated id (models echo long hex ids short, hence 4–36 and the
# resolve-by-unique-prefix below). Fullwidth 【…】 accepted on both — CJK
# bracket echo from the GLM/DeepSeek family.
#
# Keep in step with the api's app/routers/conversations.py.
CITATION_RE = re.compile(
    r"[\[【]\s*(?:c:\s*(?P<prefixed>[^\]】\s]{1,64})"
    r"|(?P<bare>[0-9a-fA-F][0-9a-fA-F-]{3,35}))\s*[\]】]"
)
# A prefix-less marker is only believed at full-id length — short bracketed hex
# is ordinary prose far more often than it is a citation.
MIN_UNPREFIXED_ID = 8
# 8-4-4-4-12: a bracketed token of exactly this shape is a marker beyond
# reasonable doubt, prefix or not, so an unresolvable one strips as a
# hallucination rather than surviving into the drafted text.
FULL_ID_RE = re.compile(r"[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}\Z")
TO_CONFIRM_RE = re.compile(r"\[TO CONFIRM:[^\]]*\]")
#: Endnote markers, once `CitationIndex.resolve` has numbered them. Removed
#: from an answer destined for a form field — see `_answer_sheet`.
FOOTNOTE_RE = re.compile(r"[ \t]*\[(\d+)\]")

#: A module's table renderer: draw `name`'s rows from the pack's own records.
TableRenderer = Callable[[Document, DraftPackBase], None]


@dataclass
class AssembledDraft:
    data: bytes
    to_confirm_count: int
    citations: list[dict] = field(default_factory=list)
    stripped_citations: int = 0
    #: One entry per section, for kinds that answer somebody else's form.
    #: Empty for ordinary documents, which are read as documents.
    answers: list[dict] = field(default_factory=list)


class CitationIndex:
    """Shared first-appearance numbering across every section."""

    def __init__(self, excerpts: list[VaultExcerpt]) -> None:
        self.by_id = {str(e.chunk_id).lower(): e for e in excerpts}
        self.order: dict[str, int] = {}
        self.stripped = 0

    def resolve(self, text: str) -> str:
        def _sub(match: re.Match) -> str:
            prefixed = match.group("prefixed")
            cid = (prefixed if prefixed is not None else match.group("bare")).lower()
            # Certain it is a marker: it carries the prefix, or it is a whole
            # uuid. An uncertain one is left verbatim unless it resolves —
            # stripping it would delete drafted prose.
            certain = prefixed is not None or FULL_ID_RE.match(cid) is not None
            if not certain and len(cid) < MIN_UNPREFIXED_ID:
                return match.group(0)
            if cid not in self.by_id:
                # Truncated markers resolve only as a unique prefix of one
                # supplied id — fabricated or ambiguous ids still strip.
                matches = [full for full in self.by_id if full.startswith(cid)]
                if len(matches) != 1:
                    if not certain:
                        return match.group(0)
                    self.stripped += 1
                    return ""
                cid = matches[0]
            if cid not in self.order:
                self.order[cid] = len(self.order) + 1
            return f"[{self.order[cid]}]"

        return CITATION_RE.sub(_sub, text)

    def references(self) -> list[dict]:
        entries = []
        for cid, n in sorted(self.order.items(), key=lambda kv: kv[1]):
            excerpt = self.by_id[cid]
            pages = f"p.{excerpt.page_start}–{excerpt.page_end}" if excerpt.page_start else None
            entries.append({"n": n, "title": excerpt.title, "pages": pages})
        return entries


def _title_page(doc: Document, pack: DraftPackBase, generated_on: date) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "DRAFT — for review"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_heading(pack.doc_title(), level=0)
    lines = [
        *pack.subject_lines(),
        f"Generated {generated_on.isoformat()} — AI-assisted draft for review",
    ]
    for line in lines:
        paragraph = doc.add_paragraph(line)
        paragraph.runs[0].font.size = Pt(12)


def _warning(doc: Document, pack: DraftPackBase) -> None:
    text = pack.warning_block()
    if not text:
        return
    paragraph = doc.add_paragraph()
    paragraph.add_run(text).bold = True


def _narrative(doc: Document, text: str) -> None:
    for block in text.split("\n\n"):
        block = block.strip()
        if block:
            doc.add_paragraph(block)


def _references(doc: Document, index: CitationIndex) -> None:
    entries = index.references()
    if not entries:
        return
    doc.add_heading("References", level=1)
    for entry in entries:
        pages = f", {entry['pages']}" if entry["pages"] else ""
        doc.add_paragraph(f"[{entry['n']}] {entry['title']}{pages}")


def _data_sources(doc: Document, pack: DraftPackBase, index: CitationIndex) -> None:
    doc.add_heading("Data sources", level=1)
    counts = ", ".join(f"{count} {name}" for name, count in pack.record_counts().items())
    doc.add_paragraph(f"Assembled from module records: {counts}.")
    cited_titles = sorted({e["title"] for e in index.references()})
    if cited_titles:
        doc.add_paragraph("Vault documents cited: " + "; ".join(cited_titles) + ".")
    for note in pack.source_notes():
        doc.add_paragraph(note)


def _answer_entry(section: Section, resolved: str, index: CitationIndex) -> dict:
    """One row of the answer sheet: prose a person can paste, plus what backs it.

    `text` has the endnote markers taken back out. They exist for a document
    with a References page at the end; a funder's form field has no such page,
    and pasting "[3]" into one asserts a reference the assessor cannot follow.
    The numbers move to `citations` instead, where the UI can show them beside
    the answer without putting them inside it.

    `[TO CONFIRM: …]` markers deliberately stay in. An answer that still has a
    gap in it should be uncomfortable to paste.
    """
    # Only numbers this document actually issued are endnotes. A model that
    # writes "[42]" as prose keeps it — the References page has no 42, so
    # deleting it would be deleting the answer's own words.
    issued = len(index.order)
    cited = {n for n in (int(m) for m in FOOTNOTE_RE.findall(resolved)) if 1 <= n <= issued}
    plain = FOOTNOTE_RE.sub(
        lambda m: "" if 1 <= int(m.group(1)) <= issued else m.group(0), resolved
    )
    # Removing a marker leaves the space that preceded it stranded in front of
    # the punctuation that followed — " ." reads as sloppy in a bid, and this
    # text is going into a funder's field verbatim.
    plain = re.sub(r"[ \t]+([.,;:!?)\]])", r"\1", plain)
    plain = re.sub(r"[ \t]{2,}", " ", plain).strip()
    length, over_by = measure(plain, section)
    return {
        "question_id": section.key,
        "question": section.title,
        "guidance": section.guidance,
        "text": plain,
        "limit": section.limit,
        "limit_kind": section.limit_kind,
        "length": length,
        "over_by": over_by,
        "to_confirm": len(TO_CONFIRM_RE.findall(plain)),
        "citations": [e for e in index.references() if e["n"] in cited],
    }


def assemble_docx(
    pack: DraftPackBase,
    sections: list[tuple[Section, str]],
    generated_on: date,
    tables: dict[str, TableRenderer] | None = None,
    answer_sheet: bool = False,
) -> AssembledDraft:
    index = CitationIndex(pack.excerpts)
    doc = Document()
    renderers = tables or {}

    _title_page(doc, pack, generated_on)
    _warning(doc, pack)
    doc.add_page_break()

    to_confirm = 0
    resolved_sections: list[tuple[Section, str]] = []
    for number, (section, text) in enumerate(sections, start=1):
        doc.add_heading(f"{number}. {section.title}", level=1)
        resolved = index.resolve(text)
        resolved_sections.append((section, resolved))
        to_confirm += len(TO_CONFIRM_RE.findall(resolved))
        _narrative(doc, resolved)
        renderer = renderers.get(section.table) if section.table else None
        if renderer is not None:
            renderer(doc, pack)

    _references(doc, index)
    _data_sources(doc, pack, index)

    buffer = io.BytesIO()
    doc.save(buffer)
    # Built after the loop, not inside it: `index.references()` is only
    # complete once every section has been resolved, and an answer's citation
    # numbers have to agree with the document's.
    answers = (
        [_answer_entry(s, resolved, index) for s, resolved in resolved_sections]
        if answer_sheet
        else []
    )
    return AssembledDraft(
        data=buffer.getvalue(),
        to_confirm_count=to_confirm,
        citations=index.references(),
        stripped_citations=index.stripped,
        answers=answers,
    )
