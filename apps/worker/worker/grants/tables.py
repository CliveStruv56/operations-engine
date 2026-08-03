"""Grantwork's data tables.

The document shell — cover page, citation resolution, `[TO CONFIRM]`
counting, references, the Data sources appendix — is shared
(`worker/drafting/assemble.py`). What belongs to Grantwork is the tables its
sections render, and the rule that makes the module trustworthy: **every
figure in a monitoring return comes from a recorded row, not from the
model.** A measure with no recorded value renders as "not recorded" rather
than as a plausible number.
"""

from typing import Any

from docx import Document

from worker.drafting.assemble import TableRenderer
from worker.grants.context import GrantPack

TABLE_STYLE = "Light Grid Accent 1"
NOT_RECORDED = "not recorded"


def _fmt(value: float | None, unit: str = "") -> str:
    if value is None:
        return "—"
    rendered = f"{value:,.0f}" if float(value).is_integer() else f"{value:,.2f}"
    return f"{rendered} {unit}".strip()


def _share(share: float | None) -> str:
    return f"{share * 100:,.0f}%" if share is not None else "—"


def _measure_rows(doc: Document, rows: list[dict[str, Any]], value_heading: str) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = TABLE_STYLE
    headings = ["Measure", "Baseline", "Target", value_heading, "Against target"]
    for i, head in enumerate(headings):
        table.rows[0].cells[i].text = head
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row["name"]
        cells[1].text = _fmt(row["baseline"], row["unit"])
        cells[2].text = _fmt(row["target"], row["unit"])
        # An unrecorded outcome says so in words. Rendering it as "0" or "—"
        # alone invites the reader — and the funder — to read absence as zero.
        cells[3].text = (
            _fmt(row["value"], row["unit"]) if row["value"] is not None else NOT_RECORDED
        )
        cells[4].text = _share(row["share"])


def impact_table(doc: Document, pack: GrantPack) -> None:
    """Measures against the reporting period being drafted.

    On a case for support or an application there is no period yet, so this
    renders the promise — baseline and target with no achieved column filled.
    """
    period = pack.target_period()
    rows = pack.measure_progress(period.id if period else None)
    if not rows:
        doc.add_paragraph("No impact measures have been recorded for this application yet.")
        return
    heading = f"Achieved ({period.label})" if period else "Achieved"
    _measure_rows(doc, rows, heading)
    if period is None:
        doc.add_paragraph(
            "Targets as proposed; no reporting period has been recorded against them yet."
        )


def outcomes_history_table(doc: Document, pack: GrantPack) -> None:
    """Latest recorded value per measure across the whole grant, for the
    end-of-grant evaluation, followed by the period-by-period detail."""
    rows = pack.cumulative_progress()
    if not rows:
        doc.add_paragraph("No impact measures have been recorded for this application.")
        return
    _measure_rows(doc, rows, "Latest recorded")

    reported = [p for p in pack.periods if pack.outcomes_for(p.id)]
    if not reported:
        return
    doc.add_paragraph("Period by period:")
    table = doc.add_table(rows=1, cols=3)
    table.style = TABLE_STYLE
    for i, head in enumerate(["Period", "Measure", "Recorded"]):
        table.rows[0].cells[i].text = head
    for period in reported:
        recorded = pack.outcomes_for(period.id)
        for measure in pack.measures:
            outcome = recorded.get(measure.id)
            if outcome is None or outcome.value is None:
                continue
            cells = table.add_row().cells
            cells[0].text = period.label
            cells[1].text = measure.name
            cells[2].text = _fmt(outcome.value, measure.unit)


def conditions_table(doc: Document, pack: GrantPack) -> None:
    """The award conditions register and where each one stands."""
    if not pack.conditions:
        doc.add_paragraph("No award conditions have been recorded for this grant.")
        return
    table = doc.add_table(rows=1, cols=4)
    table.style = TABLE_STYLE
    for i, head in enumerate(["#", "Condition", "Pre-drawdown", "Status"]):
        table.rows[0].cells[i].text = head
    for condition in pack.conditions:
        cells = table.add_row().cells
        cells[0].text = condition.number
        cells[1].text = condition.description
        cells[2].text = "Yes" if condition.pre_drawdown else "No"
        cells[3].text = condition.status.replace("_", " ")
    outstanding = [
        c for c in pack.conditions if c.status in ("outstanding", "partially_discharged")
    ]
    if outstanding:
        doc.add_paragraph(f"{len(outstanding)} condition(s) still outstanding.")


TABLES: dict[str, TableRenderer] = {
    "impact": impact_table,
    "outcomes_history": outcomes_history_table,
    "conditions": conditions_table,
}
