"""Groundwork's data tables, plus the DOCX assembly entry point.

The document shell — cover page, citation resolution, `[TO CONFIRM]`
counting, references and the Data sources appendix — is shared
(`worker/drafting/assemble.py`). What belongs to Groundwork is the pair of
tables its sections render: the budget and the funding stack, drawn from the
pack's own records rather than from anything the model wrote.
"""

from datetime import date

from docx import Document

from worker.drafting.assemble import (
    CITATION_RE,
    TO_CONFIRM_RE,
    AssembledDraft,
    TableRenderer,
)
from worker.drafting.assemble import assemble_docx as _assemble_docx
from worker.drafting.sections import Section
from worker.drafts.context import ContextPack

__all__ = [
    "CITATION_RE",
    "TABLES",
    "TO_CONFIRM_RE",
    "AssembledDraft",
    "assemble_docx",
]


def budget_table(doc: Document, pack: ContextPack) -> None:
    if not pack.budget_lines:
        return
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, head in enumerate(["Category", "Item", "Budget £", "Forecast £", "Actual £"]):
        table.rows[0].cells[i].text = head
    for line in pack.budget_lines:
        cells = table.add_row().cells
        cells[0].text = line.category
        cells[1].text = line.label
        cells[2].text = f"{line.budget:,.0f}"
        cells[3].text = f"{line.forecast:,.0f}"
        cells[4].text = f"{line.actual:,.0f}"
    totals = table.add_row().cells
    totals[0].text = "Total"
    totals[2].text = f"{pack.budget_totals.budget:,.0f}"
    totals[3].text = f"{pack.budget_totals.forecast:,.0f}"
    totals[4].text = f"{pack.budget_totals.actual:,.0f}"
    doc.add_paragraph(f"Forecast variance against budget: £{pack.budget_totals.variance:,.0f}.")


def funding_table(doc: Document, pack: ContextPack) -> None:
    if not pack.funding:
        return
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, head in enumerate(["Source", "Kind", "Status", "Sought £", "Secured £"]):
        table.rows[0].cells[i].text = head
    for source in pack.funding:
        cells = table.add_row().cells
        cells[0].text = source.name
        cells[1].text = source.kind
        cells[2].text = source.status
        cells[3].text = f"{source.amount_sought:,.0f}" if source.amount_sought else "—"
        cells[4].text = f"{source.amount_secured:,.0f}" if source.amount_secured else "—"


TABLES: dict[str, TableRenderer] = {"budget": budget_table, "funding": funding_table}


def assemble_docx(
    pack: ContextPack,
    sections: list[tuple[Section, str]],
    generated_on: date,
) -> AssembledDraft:
    return _assemble_docx(pack, sections, generated_on, tables=TABLES)
