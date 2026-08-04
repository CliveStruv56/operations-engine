"""Offline tests for Grantwork's drafting adapter: skeletons against the cost
guard, pack hooks, registry targets, the data tables and the DOCX assembly.

No database and no network — the DB-touching gather is exercised from the API
suite instead (ASSUMPTIONS #13).
"""

import io
from datetime import date
from uuid import uuid4

from docx import Document

from worker.drafting.assemble import assemble_docx
from worker.drafting.llm import MAX_LLM_CALLS
from worker.drafting.pack import VaultExcerpt
from worker.grants.context import (
    ApplicationFacts,
    CatalogueFacts,
    ConditionFacts,
    GrantPack,
    MeasureFacts,
    OutcomeFacts,
    PeriodFacts,
)
from worker.grants.prompts import SKELETONS
from worker.grants.register import registry_target
from worker.grants.retrieval import QUERY_SETS, queries_for
from worker.grants.tables import NOT_RECORDED, TABLES

TODAY = date(2026, 8, 3)
PERIOD_ID = uuid4()
LATER_PERIOD_ID = uuid4()
MEASURE_ID = uuid4()
SECOND_MEASURE_ID = uuid4()


def _application(**overrides) -> ApplicationFacts:
    return ApplicationFacts(
        **{
            "id": uuid4(),
            "title": "Community garden project",
            "application_type": "project_grant",
            "stage_current": "monitor",
            "status": "awarded",
            "amount_requested": 30_000.0,
            "amount_awarded": 27_500.0,
            "funder_name": "Borough Community Foundation",
            **overrides,
        }
    )


def _period(
    period_id=PERIOD_ID, label="Year 1", start=date(2026, 4, 1), status="open", **overrides
):
    return PeriodFacts(
        id=period_id,
        label=label,
        period_start=start,
        period_end=date(start.year + 1, 3, 31),
        status=status,
        **overrides,
    )


def _pack(kind="monitoring_report", **overrides) -> GrantPack:
    base = {
        "kind": kind,
        "generated_on": TODAY,
        "application": _application(),
        "periods": [_period()],
        "measures": [
            MeasureFacts(id=MEASURE_ID, name="Beneficiaries reached", unit="people", target=250.0),
            MeasureFacts(id=SECOND_MEASURE_ID, name="Sessions run", unit="sessions", target=40.0),
        ],
        "outcomes": [
            OutcomeFacts(
                measure_id=MEASURE_ID,
                reporting_period_id=PERIOD_ID,
                value=180.0,
                narrative="Reached 180 people across the first year.",
            )
        ],
        "target_period_id": PERIOD_ID if kind == "monitoring_report" else None,
    }
    return GrantPack(**{**base, **overrides})


# -- skeletons ---------------------------------------------------------------


def test_every_skeleton_fits_the_call_budget():
    for kind, sections in SKELETONS.items():
        assert len(sections) + 1 <= MAX_LLM_CALLS, kind  # +1 outline call


def test_every_skeleton_table_has_a_renderer():
    """A section naming a table nobody renders silently drops its figures —
    the narrative would refer to a table that never appears."""
    for kind, sections in SKELETONS.items():
        for section in sections:
            if section.table:
                assert section.table in TABLES, f"{kind}/{section.key}: no renderer"


def test_the_four_prd_document_kinds_exist():
    assert set(SKELETONS) == {
        "case_for_support",
        "funding_application",
        "monitoring_report",
        "impact_evaluation",
    }


def test_no_grantwork_section_is_told_about_a_table_it_does_not_have():
    """The bug this closes, live on 4 Aug 2026: the shared contract told every
    module that budget and funding figures were tabled, and a monitoring
    return closed by referring a funder to "the accompanying financial table".
    Grantwork has no financial table, and section 7 has no table at all."""
    from worker.drafting.prompts import section_prompt
    from worker.grants.prompts import GROUNDING_PROMPT

    financial = next(s for s in SKELETONS["monitoring_report"] if s.key == "finance")
    assert financial.table is None
    _, user = section_prompt(_pack(), financial, [], GROUNDING_PROMPT)
    assert "table" not in user.lower()

    tabled = next(s for s in SKELETONS["monitoring_report"] if s.table == "impact")
    _, with_table = section_prompt(_pack(), tabled, [], GROUNDING_PROMPT)
    assert "impact table is rendered" in with_table


def test_monitoring_reports_do_not_touch_the_vault():
    """A monitoring return accounts for what this grant did, and those facts
    are module rows. No queries means the engine skips the embedding call."""
    assert queries_for("monitoring_report", _pack()) == []
    assert not any(s.uses_vault for s in SKELETONS["monitoring_report"])
    assert set(QUERY_SETS) == {"case_for_support", "funding_application", "impact_evaluation"}


# -- pack hooks --------------------------------------------------------------


def test_doc_title_qualifies_instance_documents():
    assert _pack().doc_title() == "Monitoring return — Year 1"
    bid = _pack(kind="funding_application", target_period_id=None)
    assert bid.doc_title() == "Funding application — Borough Community Foundation"
    assert _pack(kind="case_for_support", target_period_id=None).doc_title() == "Case for support"


def test_prompt_notes_carry_the_period_and_the_figures_rule():
    notes = " ".join(_pack().prompt_notes())
    assert "Year 1" in notes and "2026-04-01" in notes
    assert "never restate or estimate" in notes


def test_unverified_catalogue_row_warns_on_page_one():
    """Seeded catalogue rows ship unverified and stale (ASSUMPTIONS #24), so a
    bid built on one must say so before anybody reads the prose."""
    catalogue = CatalogueFacts(
        key="stale_fund",
        name="Stale Fund",
        funder="Test Trust",
        kind="revenue",
        eligibility="Charities",
        status="unverified",
        last_verified=date(2026, 5, 1),
        next_review=date(2026, 5, 1),
        stale=True,
    )
    pack = _pack(kind="funding_application", target_period_id=None, catalogue=catalogue)
    warning = pack.warning_block()
    assert warning is not None
    assert "unverified" in warning and "2026-05-01" in warning
    assert "confirm the current criteria" in warning
    assert "past its review date" in warning
    assert "Warning: past its review date" in " ".join(pack.source_notes())


def test_verified_open_catalogue_row_does_not_warn():
    catalogue = CatalogueFacts(
        key="checked",
        name="Checked Fund",
        funder="Test Trust",
        kind="revenue",
        eligibility="Charities",
        status="open",
        last_verified=date(2026, 8, 1),
        next_review=date(2026, 11, 1),
        stale=False,
    )
    pack = _pack(kind="funding_application", target_period_id=None, catalogue=catalogue)
    assert pack.warning_block() is None


def test_only_bids_warn_about_the_catalogue():
    """A monitoring return is not being sent to win money on the strength of
    the catalogue row, so the warning would be noise."""
    catalogue = CatalogueFacts(
        key="stale_fund",
        name="Stale Fund",
        funder="Test Trust",
        kind="revenue",
        eligibility="Charities",
        status="unverified",
        last_verified=date(2026, 5, 1),
        next_review=date(2026, 5, 1),
        stale=True,
    )
    assert _pack(catalogue=catalogue).warning_block() is None


# -- derived progress --------------------------------------------------------


def test_measure_progress_joins_recorded_values_and_leaves_gaps_empty():
    rows = {r["name"]: r for r in _pack().measure_progress(PERIOD_ID)}
    assert rows["Beneficiaries reached"]["value"] == 180.0
    assert rows["Beneficiaries reached"]["share"] == 180 / 250
    # Nothing recorded for the second measure: it must stay None, never 0.
    assert rows["Sessions run"]["value"] is None
    assert rows["Sessions run"]["share"] is None


def test_cumulative_progress_takes_the_latest_period_per_measure():
    pack = _pack(
        periods=[_period(), _period(LATER_PERIOD_ID, "Year 2", date(2027, 4, 1))],
        outcomes=[
            OutcomeFacts(measure_id=MEASURE_ID, reporting_period_id=PERIOD_ID, value=180.0),
            OutcomeFacts(measure_id=MEASURE_ID, reporting_period_id=LATER_PERIOD_ID, value=260.0),
        ],
    )
    rows = {r["name"]: r for r in pack.cumulative_progress()}
    assert rows["Beneficiaries reached"]["value"] == 260.0
    assert rows["Beneficiaries reached"]["share"] == 260 / 250  # over-delivery is not capped


def test_measure_progress_survives_a_zero_target():
    pack = _pack(measures=[MeasureFacts(id=MEASURE_ID, name="Zero", unit="x", target=0.0)])
    assert pack.measure_progress(PERIOD_ID)[0]["share"] is None


# -- registry targets --------------------------------------------------------


def test_monitoring_returns_get_per_period_registry_rows():
    """grant_documents is unique on (application_id, doc_type_key), so each
    period's return needs its own key (ASSUMPTIONS #9)."""
    key, title = registry_target(_pack())
    assert key == f"monitoring_report_{str(PERIOD_ID).replace('-', '')[:8]}"
    assert title == "Monitoring return — Year 1"

    other = _pack(
        periods=[_period(), _period(LATER_PERIOD_ID, "Year 2", date(2027, 4, 1))],
        target_period_id=LATER_PERIOD_ID,
    )
    assert registry_target(other)[0] != key


def test_one_off_documents_version_onto_their_seeded_row():
    for kind in ("case_for_support", "funding_application", "impact_evaluation"):
        key, _ = registry_target(_pack(kind=kind, target_period_id=None))
        assert key == kind


# -- tables and assembly -----------------------------------------------------


def _render(pack: GrantPack, kind: str | None = None) -> Document:
    sections = SKELETONS[kind or pack.kind]
    drafted = [(s, f"Narrative for {s.key}.") for s in sections]
    result = assemble_docx(pack, drafted, TODAY, tables=TABLES)
    return Document(io.BytesIO(result.data))


def _table_text(doc: Document) -> str:
    return "\n".join(
        "\t".join(cell.text for cell in row.cells) for table in doc.tables for row in table.rows
    )


def test_impact_table_says_not_recorded_rather_than_zero():
    """Absence read as zero is the failure that would embarrass a charity in
    front of its funder."""
    text = _table_text(_render(_pack()))
    assert "Beneficiaries reached" in text
    assert "180 people" in text
    assert NOT_RECORDED in text
    assert "72%" in text  # 180/250


def test_conditions_table_renders_the_register():
    pack = _pack(
        conditions=[
            ConditionFacts(
                number="1",
                description="Acknowledge the funder",
                pre_drawdown=False,
                status="discharged",
            ),
            ConditionFacts(
                number="6",
                description="Confirm match funding",
                pre_drawdown=True,
                status="outstanding",
            ),
        ]
    )
    text = _table_text(_render(pack))
    assert "Acknowledge the funder" in text
    assert "Confirm match funding" in text


def test_evaluation_renders_period_by_period_history():
    pack = _pack(
        kind="impact_evaluation",
        target_period_id=None,
        periods=[_period(), _period(LATER_PERIOD_ID, "Year 2", date(2027, 4, 1))],
        outcomes=[
            OutcomeFacts(measure_id=MEASURE_ID, reporting_period_id=PERIOD_ID, value=180.0),
            OutcomeFacts(measure_id=MEASURE_ID, reporting_period_id=LATER_PERIOD_ID, value=260.0),
        ],
    )
    text = _table_text(_render(pack))
    assert "Year 1" in text and "Year 2" in text
    assert "180 people" in text and "260 people" in text


def test_assembly_counts_to_confirm_and_strips_fake_citations():
    excerpt = VaultExcerpt(
        chunk_id=uuid4(),
        document_id=uuid4(),
        title="Needs assessment 2026",
        content="Local demand is high.",
    )
    pack = _pack(kind="case_for_support", target_period_id=None, excerpts=[excerpt])
    sections = SKELETONS["case_for_support"]
    drafted = [(s, "Body.") for s in sections]
    drafted[0] = (
        sections[0],
        f"Real claim [c:{excerpt.chunk_id}]. Invented claim [c:{uuid4()}]. "
        "[TO CONFIRM: trustee count] and [TO CONFIRM: year founded].",
    )
    result = assemble_docx(pack, drafted, TODAY, tables=TABLES)
    assert result.to_confirm_count == 2
    assert result.stripped_citations == 1
    assert [c["title"] for c in result.citations] == ["Needs assessment 2026"]


def test_a_pack_with_no_measures_still_assembles():
    """An application drafted before anyone recorded a measure must produce a
    document, not a crash — the section says so in words instead."""
    doc = _render(_pack(measures=[], outcomes=[]))
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "No impact measures have been recorded" in body
