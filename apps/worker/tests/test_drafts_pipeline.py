"""Offline tests for the drafting pipeline: cost guard, outline parsing,
prompt construction, and DOCX assembly (citations, [TO CONFIRM] counting,
programme status warning, data tables)."""

import io
from datetime import date
from uuid import uuid4

import pytest
from docx import Document

from tests.test_drafts_context import _funding, _pack, _programme
from worker.drafts.assemble import assemble_docx
from worker.drafts.context import BudgetLine, BudgetTotals, VaultExcerpt
from worker.drafts.llm import (
    MAX_CONTEXT_TOKENS_PER_CALL,
    MAX_LLM_CALLS,
    DraftBudgetExceeded,
    LlmCall,
    LlmLedger,
    chat,
)
from worker.drafts.prompts import SKELETONS, parse_outline, section_prompt
from worker.drafts.register import registry_target


def test_cost_guard_call_count():
    ledger = LlmLedger()
    ledger.calls = [LlmCall("drafter", 10, 10)] * MAX_LLM_CALLS
    with pytest.raises(DraftBudgetExceeded, match="model-call budget"):
        ledger.check_next_call("s", "u")


def test_cost_guard_context_size():
    ledger = LlmLedger()
    huge = "x" * (MAX_CONTEXT_TOKENS_PER_CALL * 4 + 8)
    with pytest.raises(DraftBudgetExceeded, match="more context"):
        ledger.check_next_call("system", huge)


def test_every_skeleton_fits_the_call_budget():
    for kind, sections in SKELETONS.items():
        assert len(sections) + 1 <= MAX_LLM_CALLS, kind  # +1 outline call


def test_ledger_cost_math():
    ledger = LlmLedger()
    ledger.calls = [LlmCall("drafter", 1_000_000, 0), LlmCall("reasoner", 0, 1_000_000)]
    ledger.embed_cost_usd = 0.01
    assert ledger.cost_usd == pytest.approx(0.15 + 3.00 + 0.01)
    assert ledger.tokens_in == 1_000_000 and ledger.tokens_out == 1_000_000


def test_parse_outline_accepts_json_and_fences():
    assert parse_outline('{"exec": ["note one"]}') == {"exec": ["note one"]}
    fenced = '```json\n{"exec": ["a", "b"]}\n```'
    assert parse_outline(fenced) == {"exec": ["a", "b"]}
    assert parse_outline("Sorry, here is prose.") == {}
    assert parse_outline('["not", "a", "dict"]') == {}


def test_section_prompt_vault_variants():
    section = next(s for s in SKELETONS["feasibility_study"] if s.uses_vault)
    excerpt = VaultExcerpt(
        chunk_id=uuid4(), document_id=uuid4(), title="Local Plan", page_start=4, content="policy"
    )
    pack = _pack(kind="feasibility_study", excerpts=[excerpt])
    _, user = section_prompt(pack, section, ["note"])
    assert f"[c:{excerpt.chunk_id}]" in user and "<vault-excerpts>" in user

    _, empty_user = section_prompt(_pack(kind="feasibility_study"), section, [])
    assert "No vault excerpts were found" in empty_user


def _assembly_pack(**overrides):
    excerpt = VaultExcerpt(
        chunk_id=uuid4(),
        document_id=uuid4(),
        title="Housing Needs Survey",
        page_start=12,
        page_end=13,
        content="Forty households responded.",
    )
    base = dict(
        kind="monthly_report",
        report_month="2026-07",
        budget_lines=[
            BudgetLine(
                category="land", label="Acquisition", budget=250000, forecast=250000, actual=0
            )
        ],
        budget_totals=BudgetTotals(budget=250000, forecast=250000, actual=0, variance=0),
        funding=[_funding()],
        excerpts=[excerpt],
    )
    base.update(overrides)
    return _pack(**base), excerpt


def _texts(data: bytes) -> list[str]:
    doc = Document(io.BytesIO(data))
    return [p.text for p in doc.paragraphs]


def test_assemble_monthly_report_structure_citations_and_to_confirm():
    pack, excerpt = _assembly_pack()
    sections = SKELETONS["monthly_report"]
    texts = [f"Narrative for {s.title}." for s in sections]
    # Section 1 carries a real citation, a fabricated one, and two TO CONFIRMs.
    # A full id, a model-truncated 8-char prefix of the same id (must share
    # its numbering), and a fabricated id (must strip).
    texts[0] = (
        f"On track overall [c:{excerpt.chunk_id}] and also [c:{uuid4()}]. "
        f"Truncated echo [c:{str(excerpt.chunk_id)[:8]}]. "
        "[TO CONFIRM: contract start date] [TO CONFIRM: insurance renewal]"
    )
    draft = assemble_docx(pack, list(zip(sections, texts, strict=True)), date(2026, 7, 31))

    assert draft.to_confirm_count == 2
    assert draft.stripped_citations == 1
    assert draft.citations == [{"n": 1, "title": "Housing Needs Survey", "pages": "p.12–13"}]

    paragraphs = _texts(draft.data)
    body = "\n".join(paragraphs)
    assert "Monthly client report — 2026-07" in body
    for i, section in enumerate(sections, start=1):
        assert f"{i}. {section.title}" in body
    assert "On track overall [1] and also ." in body  # resolved + stripped
    assert "Truncated echo [1]." in body  # prefix marker shares the numbering
    assert "[TO CONFIRM: contract start date]" in body  # markers stay visible
    assert "References" in body and "[1] Housing Needs Survey, p.12–13" in body
    assert "Data sources" in body and "10 stages" not in body  # counts are real, not invented

    doc = Document(io.BytesIO(draft.data))
    assert doc.sections[0].header.paragraphs[0].text == "DRAFT — for review"
    assert len(doc.tables) == 2  # budget + funding


def test_assemble_funding_bid_warning_block():
    linked = _funding(programme_key="cwmpas_cch")
    announced = _programme("cwmpas_cch", status="announced")
    pack, _ = _assembly_pack(
        kind="funding_bid",
        report_month=None,
        funding=[linked],
        programmes=[announced],
        target_funding_id=linked.id,
    )
    sections = SKELETONS["funding_bid"]
    draft = assemble_docx(pack, [(s, f"Text for {s.title}.") for s in sections], date(2026, 7, 31))
    body = "\n".join(_texts(draft.data))
    assert "Programme status was “announced” when last verified 2026-07-28" in body
    assert "Funding application — CLT grant" in body

    # An open programme renders no warning.
    open_pack, _ = _assembly_pack(
        kind="funding_bid",
        report_month=None,
        funding=[linked],
        programmes=[_programme("cwmpas_cch", status="open")],
        target_funding_id=linked.id,
    )
    open_draft = assemble_docx(
        open_pack, [(s, f"Text for {s.title}.") for s in sections], date(2026, 7, 31)
    )
    assert "confirm before submitting" not in "\n".join(_texts(open_draft.data))


def test_registry_target_keys_and_titles():
    monthly, _ = _assembly_pack()
    assert registry_target(monthly) == (
        "monthly_report_2026_07",
        "Monthly client report — July 2026",
    )

    linked = _funding(programme_key="cwmpas_cch")
    bid, _ = _assembly_pack(
        kind="funding_bid", report_month=None, funding=[linked], target_funding_id=linked.id
    )
    key, title = registry_target(bid)
    assert key == f"funding_bid_{str(linked.id).replace('-', '')[:8]}"
    assert title == "Funding application — CLT grant"

    feasibility, _ = _assembly_pack(kind="feasibility_study", report_month=None)
    assert registry_target(feasibility) == ("feasibility_study", "Feasibility study")


# -- reasoning-model output budget -------------------------------------------
#
# Found by the first live Grantwork draft (3 Aug 2026). Both drafting aliases
# are reasoning models that bill thinking against `completion_tokens`. At the
# old 1024 ceiling every section of a data-heavy document came back
# `finish_reason=length`, and the ones that reasoned longest returned no
# content at all — which the pipeline assembled into a document with six of
# nine sections missing, filed as a clean draft with to_confirm_count 0.


class _FakeResponse:
    def __init__(self, payload):
        self._payload = payload

    def raise_for_status(self):
        return None

    def json(self):
        return self._payload


class _FakeClient:
    """Stands in for httpx.AsyncClient, returning a canned completion."""

    def __init__(self, payload):
        self._payload = payload
        self.sent = None

    async def __aenter__(self):
        return self

    async def __aexit__(self, *exc):
        return False

    async def post(self, _url, headers=None, json=None):
        self.sent = json
        return _FakeResponse(self._payload)


def _completion(content, finish_reason="stop", completion_tokens=200):
    return {
        "choices": [{"message": {"content": content}, "finish_reason": finish_reason}],
        "usage": {"prompt_tokens": 100, "completion_tokens": completion_tokens},
    }


def _patch(monkeypatch, payload):
    client = _FakeClient(payload)
    monkeypatch.setattr("worker.drafting.llm.httpx.AsyncClient", lambda **kw: client)
    monkeypatch.setattr(
        "worker.drafting.llm.get_settings",
        lambda: type("S", (), {"litellm_base_url": "http://gateway"})(),
    )
    return client


async def test_output_budget_clears_reasoning_tokens(monkeypatch):
    """The ceiling has to fit thinking plus prose, not prose alone."""
    from worker.drafting.llm import MAX_OUTPUT_TOKENS

    client = _patch(monkeypatch, _completion("Some prose."))
    await chat(LlmLedger(), "key", "drafter", "sys", "usr")
    assert client.sent["max_tokens"] == MAX_OUTPUT_TOKENS
    assert MAX_OUTPUT_TOKENS >= 4096, "must clear ~700 reasoning tokens plus real prose"


async def test_an_empty_section_fails_the_job_rather_than_shrinking_the_document(monkeypatch):
    from worker.drafting.llm import EmptySectionError

    _patch(monkeypatch, _completion("", finish_reason="length"))
    with pytest.raises(EmptySectionError, match="empty section"):
        await chat(LlmLedger(), "key", "drafter", "sys", "usr")


async def test_a_null_content_is_treated_as_empty(monkeypatch):
    """Some gateways return null rather than "" when a reasoning model spends
    its whole budget thinking."""
    from worker.drafting.llm import EmptySectionError

    _patch(monkeypatch, _completion(None, finish_reason="length"))
    with pytest.raises(EmptySectionError):
        await chat(LlmLedger(), "key", "drafter", "sys", "usr")


async def test_the_outline_call_may_come_back_empty(monkeypatch):
    """The outline is an optimisation, designed to degrade to {} rather than
    spend budget on retries — so it must not fail the job."""
    _patch(monkeypatch, _completion(""))
    assert await chat(LlmLedger(), "key", "drafter", "sys", "usr", allow_empty=True) == ""
    assert parse_outline("") == {}


async def test_truncation_is_recorded_on_the_ledger(monkeypatch):
    _patch(monkeypatch, _completion("Cut off mid-", finish_reason="length"))
    ledger = LlmLedger()
    await chat(ledger, "key", "drafter", "sys", "usr")
    assert ledger.calls[-1].truncated is True
    assert ledger.truncated_calls == 1

    _patch(monkeypatch, _completion("Complete."))
    await chat(ledger, "key", "drafter", "sys", "usr")
    assert ledger.calls[-1].truncated is False
    assert ledger.truncated_calls == 1


async def test_an_empty_section_is_retried_once_before_failing(monkeypatch):
    """Empty replies are intermittent, so one retry turns a transient into a
    non-event — without ever falling back to a document with a gap."""
    from worker.drafting.engine import _section_text
    from worker.drafting.llm import EmptySectionError
    from worker.drafting.sections import Section

    replies = iter([_completion(""), _completion("Recovered prose.")])
    monkeypatch.setattr(
        "worker.drafting.llm.httpx.AsyncClient", lambda **kw: _FakeClient(next(replies))
    )
    monkeypatch.setattr(
        "worker.drafting.llm.get_settings",
        lambda: type("S", (), {"litellm_base_url": "http://gateway"})(),
    )
    ledger = LlmLedger()
    text = await _section_text(ledger, "key", Section("k", "T"), "sys", "usr")
    assert text == "Recovered prose."
    assert len(ledger.calls) == 2, "the retry is charged to the ledger, not hidden"

    both_empty = iter([_completion(""), _completion("")])
    monkeypatch.setattr(
        "worker.drafting.llm.httpx.AsyncClient", lambda **kw: _FakeClient(next(both_empty))
    )
    with pytest.raises(EmptySectionError):
        await _section_text(LlmLedger(), "key", Section("k", "T"), "sys", "usr")


async def test_reasoning_effort_is_bounded_on_every_call(monkeypatch):
    """A reasoning model given an unbounded budget will spend all of it
    thinking and return no prose — measured live on the `reasoner` alias."""
    from worker.drafting.llm import REASONING_EFFORT

    client = _patch(monkeypatch, _completion("Prose."))
    await chat(LlmLedger(), "key", "reasoner", "sys", "usr")
    assert client.sent["reasoning_effort"] == REASONING_EFFORT
    assert REASONING_EFFORT in ("none", "low", "medium")
